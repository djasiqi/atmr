"""Composition DOCX du contrat-cadre partenaire LIRIE (python-docx)."""

from __future__ import annotations

import hashlib
import io
import re
from decimal import Decimal, InvalidOperation
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

from models.enums import LegalForm
from services.platform_billing.partner_agreement_compliance import (
    RETENTION_CATEGORIES,
    active_technical_providers,
)

# Source unique des versions du contrat partenaire (v1.20).
# GENERATOR_VERSION est ré-exporté ici pour compatibilité
# (`from partner_agreement_docx import GENERATOR_VERSION, TEMPLATE_VERSION`).
from services.platform_billing.partner_agreement_versions import (  # noqa: F401
    GENERATOR_VERSION,
    PACK_SCHEMA_VERSION,
    PENALTY_CALCULATION_VERSION,
    PENALTY_CURRENCY,
    PENALTY_MINIMUM_CHF,
    PENALTY_MULTIPLIER,
    RETENTION_POLICY_VERSION,
    SUBPROCESSORS_VERSION,
    TEMPLATE_VERSION,
)

# Le builder historique A+B+C reste disponible pour lecture / comparaison.
# Le pack officiel est ParticularAgreementContent + PDF (voir partner_agreement.py).
assert TEMPLATE_VERSION == PACK_SCHEMA_VERSION

TEMPLATE_RELATIVE = Path("templates/contracts/lirie_partenariat_base_v1.docx")

DEFAULT_OPERATOR_CONTRACTUAL_EMAIL = "info@lirie.ch"

# Charte LIRIE (tokens.css --brand-primary)
LIRIE_GREEN = RGBColor(0x00, 0x79, 0x6B)
# Police documents LIRIE — sans-serif standard (équivalent Helvetica des PDF/emails)
LIRIE_FONT = "Calibri"
BODY_SIZE_PT = 10.5
# Espacements (pt) — aérés mais mesurés
SPACE_BODY_AFTER_PT = 8
SPACE_HEADING_BEFORE_PT = 12
SPACE_HEADING_AFTER_PT = 6
SPACE_BULLET_AFTER_PT = 3
SPACE_PARTY_LINE_AFTER_PT = 3
LINE_SPACING = 1.15

_LOGO_CANDIDATES = (
    Path(__file__).resolve().parents[2] / "assets" / "lirie" / "logo-lirie.png",
    Path("/app/assets/lirie/logo-lirie.png"),
    Path("/app/backend/assets/lirie/logo-lirie.png"),
)

_CANCEL_LABELS = {
    "exclude": (
        "les courses annulées, y compris celles donnant lieu à des frais "
        "d'annulation, sont exclues de la base de commission"
    ),
    "on_cancellation_fees": (
        "lorsqu'un frais d'annulation est effectivement facturé au client final, "
        "la commission est calculée sur le montant HT de ce seul frais ; "
        "sinon la course annulée est exclue"
    ),
    "on_billed_amount": (
        "la commission est calculée sur le montant HT facturé au client final, "
        "y compris en cas d'annulation facturée"
    ),
}

_BRAND_ALIASES = frozenset({"lirie", "lirie.ch", "lirie sa", "lirie sàrl"})

_MODE_LABELS = {
    "free": "Gratuit",
    "fixed": "Montant fixe",
    "volume": "Volume mensuel (grille de paliers)",
}

_FR_MONTHS = (
    "janvier",
    "février",
    "mars",
    "avril",
    "mai",
    "juin",
    "juillet",
    "août",
    "septembre",
    "octobre",
    "novembre",
    "décembre",
)

_FR_SMALL_NUMBERS = (
    "zéro",
    "un",
    "deux",
    "trois",
    "quatre",
    "cinq",
    "six",
    "sept",
    "huit",
    "neuf",
    "dix",
    "onze",
    "douze",
    "treize",
    "quatorze",
    "quinze",
    "seize",
    "dix-sept",
    "dix-huit",
    "dix-neuf",
    "vingt",
)

# Nombres fréquents hors de la série 0–20 (délais contractuels).
_FR_EXTRA_NUMBERS = {
    24: "vingt-quatre",
    30: "trente",
    60: "soixante",
    90: "quatre-vingt-dix",
}


def _num_to_word_fr(value: int) -> str:
    """Écrit en toutes lettres les petits nombres (style rédactionnel juridique)."""
    if 0 <= value < len(_FR_SMALL_NUMBERS):
        return _FR_SMALL_NUMBERS[value]
    if value in _FR_EXTRA_NUMBERS:
        return _FR_EXTRA_NUMBERS[value]
    return str(value)


def _fmt_effective_date_fr(iso_or_str: str | None) -> str:
    """Convertit une date ISO (« 2026-08-01 ») en date française (« 1er août 2026 »)."""
    text = (str(iso_or_str) if iso_or_str else "").strip()
    if not text:
        return "—"
    match = re.match(r"^(\d{4})-(\d{2})-(\d{2})", text)
    if not match:
        return text
    year, month, day = (int(match.group(i)) for i in (1, 2, 3))
    if not 1 <= month <= 12:
        return text
    day_txt = "1er" if day == 1 else str(day)
    return f"{day_txt} {_FR_MONTHS[month - 1]} {year}"


def _fmt_chf_amount(value: Any) -> str:
    """Formatte un montant CHF façon suisse : 1000 -> « 1'000.– », 45.5 -> « 45.50 »."""
    try:
        dec = Decimal(str(value).replace(",", "."))
    except (InvalidOperation, ValueError, TypeError):
        return str(value)
    dec = dec.quantize(Decimal("0.01"))
    sign = "-" if dec < 0 else ""
    dec_abs = abs(dec)
    if dec_abs == dec_abs.to_integral_value():
        int_fmt = f"{int(dec_abs):,}".replace(",", "'")
        return f"{sign}{int_fmt}.–"
    int_part = int(dec_abs)
    cents = int((dec_abs - int_part) * 100)
    int_fmt = f"{int_part:,}".replace(",", "'")
    return f"{sign}{int_fmt}.{cents:02d}"


def template_path() -> Path:
    # backend/services/platform_billing -> backend/
    backend_root = Path(__file__).resolve().parents[2]
    return backend_root / TEMPLATE_RELATIVE


def _resolve_lirie_logo_path() -> Path | None:
    for path in _LOGO_CANDIDATES:
        if path.is_file():
            return path
    return None


def ensure_base_template() -> Path:
    """Crée le document de base (styles) s'il est absent."""
    path = template_path()
    path.parent.mkdir(parents=True, exist_ok=True)
    if path.is_file():
        return path
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    doc.add_heading("LIRIE — modèle de base contrat partenaire", level=1)
    doc.add_paragraph(
        "Document de styles uniquement. Le contenu juridique est composé par code."
    )
    doc.save(str(path))
    return path


def template_sha256() -> str:
    path = ensure_base_template()
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _pct(rate: str | None) -> str:
    if rate is None or rate == "":
        return "—"
    try:
        n = float(str(rate).replace(",", "."))
        return f"{(n * 100):.2f}".rstrip("0").rstrip(".") + " %"
    except ValueError:
        return str(rate)


def _fmt_ide(uid_ide: str | None) -> str:
    value = (uid_ide or "").strip()
    return value if value else "non attribué"


def _normalize_street(value: str) -> str:
    """Corrige « Ernest- Pictet » → « Ernest-Pictet »."""
    text = (value or "").strip()
    text = re.sub(r"\s*-\s*", "-", text)
    return re.sub(r"\s+", " ", text)


def _fmt_address(party: dict[str, Any]) -> str:
    addr = _normalize_street(party.get("street_name") or "")
    if party.get("building_number"):
        addr = f"{addr} {party['building_number']}".strip()
    postal = (party.get("postal_code") or "").strip()
    city = (party.get("city") or "").strip()
    country = (party.get("country_code") or "CH").strip()
    country_label = "Suisse" if country.upper() == "CH" else country
    return f"{addr}, {postal} {city}, {country_label}".strip(", ")


def _is_sole_proprietor(party: dict[str, Any]) -> bool:
    return (party.get("legal_form") or "") == LegalForm.SOLE_PROPRIETORSHIP.value


def _operator_natural_person_name(party: dict[str, Any]) -> str:
    """Personne physique de l'Exploitant (pas l'enseigne LIRIE)."""
    signatory = (party.get("signatory_name") or "").strip()
    legal = (party.get("legal_name") or "").strip()
    if _is_sole_proprietor(party):
        if signatory:
            return signatory
        if legal and legal.lower() not in _BRAND_ALIASES:
            return legal
        return legal or "—"
    return legal or signatory or "—"


def _operator_display_name(party: dict[str, Any]) -> str:
    """Identification juridique de l'Exploitant (indépendant sous enseigne LIRIE)."""
    name = _operator_natural_person_name(party)
    if _is_sole_proprietor(party):
        return f"{name}, exerçant en qualité d'indépendant sous l'enseigne LIRIE"
    label = party.get("legal_form_label") or party.get("legal_form") or ""
    if label:
        return f"{name} ({label}), exploitant la plateforme LIRIE"
    return f"{name}, exploitant la plateforme LIRIE"


def _partner_legal_name(party: dict[str, Any]) -> str:
    """Raison sociale complète (ex. Emmenez-moi Sàrl)."""
    name = (party.get("legal_name") or "").strip() or "—"
    form = party.get("legal_form") or ""
    label = (party.get("legal_form_label") or "").strip()
    if (
        form == LegalForm.SARL.value
        and label
        and "sàrl" not in name.lower()
        and "sarl" not in name.lower()
    ):
        return f"{name} {label}"
    if (
        form == LegalForm.SA.value
        and label
        and not re.search(r"\bSA\b", name, flags=re.IGNORECASE)
    ):
        return f"{name} {label}"
    return name


def _contractual_email(party: dict[str, Any], *, is_operator: bool) -> str:
    email = (party.get("contractual_email") or "").strip()
    if email:
        return email
    if is_operator:
        return DEFAULT_OPERATOR_CONTRACTUAL_EMAIL
    return "à compléter"


def _normalize_signatory_title(title: str) -> str:
    """Normalise « Associé(e) Gérant » → « associé-gérant » (forme contractuelle)."""
    text = (title or "").strip()
    if not text:
        return text
    return re.sub(
        r"(?i)\bassoci(?:é|ée)\s*-?\s*g[ée]rant\b",
        "associé-gérant",
        text,
    )


def _party_block(*, role: str, party: dict[str, Any], is_operator: bool) -> list[str]:
    if is_operator:
        identity_line = _operator_display_name(party)
        domicile_label = "Domicilié à"
        form_line = (
            "Statut : indépendant"
            if _is_sole_proprietor(party)
            else f"Forme juridique : {party.get('legal_form_label') or '—'}"
        )
    else:
        identity_line = _partner_legal_name(party)
        domicile_label = "Siège"
        form_line = f"Forme juridique : {party.get('legal_form_label') or '—'}"

    lines = [
        role,
        identity_line,
        form_line,
        f"{domicile_label} : {_fmt_address(party)}",
        f"IDE : {_fmt_ide(party.get('uid_ide'))}",
        f"Courriel contractuel : {_contractual_email(party, is_operator=is_operator)}",
    ]

    if is_operator and _is_sole_proprietor(party):
        # L'indépendant agit personnellement — pas de représentation.
        return lines

    signatory = (party.get("signatory_name") or "").strip() or "—"
    title = _normalize_signatory_title(party.get("signatory_title") or "")
    represented = f"Représenté par : {signatory}"
    if title:
        represented = f"{represented}, {title}"
    lines.append(represented)
    return lines


def _set_run_size(run: Any, size_pt: float) -> None:
    run.font.size = Pt(size_pt)


def _apply_lirie_run(
    run: Any,
    *,
    size_pt: float | None = None,
    bold: bool | None = None,
    color: RGBColor | None = None,
) -> None:
    """Police + couleur LIRIE (écrase le bleu thème Word des titres)."""
    run.font.name = LIRIE_FONT
    r_el = run._element
    rPr = r_el.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), LIRIE_FONT)
    rFonts.set(qn("w:hAnsi"), LIRIE_FONT)
    rFonts.set(qn("w:cs"), LIRIE_FONT)
    rFonts.set(qn("w:eastAsia"), LIRIE_FONT)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _configure_lirie_styles(doc: DocxDocument) -> None:
    """Styles document : police LIRIE + titres en vert (pas le bleu accent Word)."""
    for name in ("Normal", "List Bullet", "List Number"):
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        style.font.name = LIRIE_FONT
        style.font.size = Pt(BODY_SIZE_PT)
        style.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    for level in (1, 2, 3):
        try:
            style = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        style.font.name = LIRIE_FONT
        style.font.bold = True
        style.font.color.rgb = LIRIE_GREEN
        style.font.size = Pt(12 if level == 1 else 11)


def _compact_paragraph(
    p: Paragraph,
    *,
    space_before_pt: float = 0,
    space_after_pt: float = SPACE_BODY_AFTER_PT,
    keep_with_next: bool = False,
) -> Paragraph:
    pf = p.paragraph_format
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    pf.line_spacing = LINE_SPACING
    pf.keep_with_next = keep_with_next
    return p


def _add_heading(
    doc: DocxDocument,
    text: str,
    level: int = 1,
    *,
    keep_with_next: bool = True,
    space_before_pt: float = SPACE_HEADING_BEFORE_PT,
    space_after_pt: float = SPACE_HEADING_AFTER_PT,
) -> Paragraph:
    p = doc.add_heading(text, level=level)
    _compact_paragraph(
        p,
        space_before_pt=space_before_pt,
        space_after_pt=space_after_pt,
        keep_with_next=keep_with_next,
    )
    for run in p.runs:
        _apply_lirie_run(
            run,
            size_pt=12 if level == 1 else 11,
            bold=True,
            color=LIRIE_GREEN,
        )
    return p


def _add_para(
    doc: DocxDocument,
    text: str,
    *,
    bold: bool = False,
    keep_with_next: bool = False,
    space_before_pt: float = 0,
    space_after_pt: float = SPACE_BODY_AFTER_PT,
    size_pt: float | None = None,
) -> Paragraph:
    p = doc.add_paragraph()
    run = p.add_run(text)
    _apply_lirie_run(
        run,
        size_pt=size_pt if size_pt is not None else BODY_SIZE_PT,
        bold=bold,
    )
    _compact_paragraph(
        p,
        space_before_pt=space_before_pt,
        space_after_pt=space_after_pt,
        keep_with_next=keep_with_next,
    )
    return p


def _add_bullets(doc: DocxDocument, items: list[str]) -> None:
    for i, item in enumerate(items):
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            _apply_lirie_run(run, size_pt=BODY_SIZE_PT)
        _compact_paragraph(
            p,
            space_before_pt=0,
            space_after_pt=SPACE_BULLET_AFTER_PT,
            keep_with_next=i < len(items) - 1,
        )


def _add_field_run(paragraph: Paragraph, instr: str) -> None:
    """Insère un champ Word (PAGE, NUMPAGES, …)."""
    run = paragraph.add_run()
    _apply_lirie_run(run, size_pt=8)
    r = run._r
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = f" {instr} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text_el = OxmlElement("w:t")
    # Valeur de repli affichée avant recalcul du champ par Word (PAGE/NUMPAGES).
    text_el.text = "1"
    text_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "16")  # demi-points → 8 pt
    rPr.append(sz)
    text_run.append(rPr)
    text_run.append(text_el)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r.append(begin)
    r.append(instr_el)
    r.append(separate)
    r.append(text_run)
    r.append(end)


def _configure_section_and_footer(doc: DocxDocument, *, reference: str) -> None:
    """Marges + pied de page (référence + pagination) sur toutes les pages."""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(2.2)
    section.footer_distance = Cm(1.0)

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    # Vider le contenu du premier paragraphe
    for child in list(p._p):
        if not child.tag.endswith("pPr"):
            p._p.remove(child)

    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    tabs = pf.tab_stops
    usable_emu = int(section.page_width - section.left_margin - section.right_margin)
    tabs.add_tab_stop(usable_emu // 2, WD_TAB_ALIGNMENT.CENTER)
    tabs.add_tab_stop(usable_emu, WD_TAB_ALIGNMENT.RIGHT)

    left = p.add_run("LIRIE  ·  www.lirie.ch")
    _apply_lirie_run(left, size_pt=8, color=LIRIE_GREEN)

    p.add_run("\t")
    mid = p.add_run(f"Réf. {reference}")
    _apply_lirie_run(mid, size_pt=8)

    p.add_run("\t")
    page_label = p.add_run("Page ")
    _apply_lirie_run(page_label, size_pt=8)
    _add_field_run(p, "PAGE")
    sep = p.add_run(" / ")
    _apply_lirie_run(sep, size_pt=8)
    _add_field_run(p, "NUMPAGES")

    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "6")
    top.set(qn("w:space"), "4")
    top.set(qn("w:color"), "CCCCCC")
    pBdr.append(top)
    pPr.append(pBdr)


def _add_signature_cell(cell: Any, *, label: str, signatory: str) -> None:
    lines = [
        (label, True),
        ((signatory or "").strip() or "________________", False),
        ("", False),
        ("Lieu : ________________________", False),
        ("Date : ________________________", False),
        ("", False),
        ("Signature :", False),
        ("", False),
        ("", False),
        ("______________________________", False),
    ]
    cell.text = ""
    for idx, (text, bold) in enumerate(lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        if idx == 0:
            for child in list(p._p):
                if not child.tag.endswith("pPr"):
                    p._p.remove(child)
        if text:
            run = p.add_run(text)
            _apply_lirie_run(
                run,
                size_pt=BODY_SIZE_PT,
                bold=bold,
                color=LIRIE_GREEN if bold else None,
            )
        _compact_paragraph(p, space_before_pt=0, space_after_pt=2)


def _add_signatures_table(
    doc: DocxDocument,
    *,
    operator_signatory: str,
    partner_signatory: str | None,
    partner_co_signatory: str | None = None,
) -> None:
    """Deux colonnes côte à côte — évite une page quasi vide de signatures seules.

    Si ``partner_co_signatory`` est fourni (signature collective attestée côté
    Partenaire), une seconde ligne accueille la seconde signature.
    """
    rows = 2 if partner_co_signatory else 1
    table = doc.add_table(rows=rows, cols=2)
    table.autofit = True
    _add_signature_cell(
        table.rows[0].cells[0],
        label="Pour l'Exploitant :",
        signatory=operator_signatory,
    )
    _add_signature_cell(
        table.rows[0].cells[1],
        label=(
            "Pour le Partenaire (signature 1 — collective) :"
            if partner_co_signatory
            else "Pour le Partenaire :"
        ),
        signatory=(partner_signatory or "").strip() or "________________",
    )
    if partner_co_signatory:
        table.rows[1].cells[0].text = ""
        _add_signature_cell(
            table.rows[1].cells[1],
            label="Pour le Partenaire (signature 2 — collective) :",
            signatory=partner_co_signatory,
        )
    # Espacement léger après le tableau
    _add_para(doc, "", space_after_pt=0)


def _add_simple_table(
    doc: DocxDocument,
    *,
    headers: list[str],
    rows: list[list[str]],
) -> None:
    """Tableau sobre (en-tête vert LIRIE) pour paliers / matrices."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for col_idx, header in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        _apply_lirie_run(run, size_pt=BODY_SIZE_PT, bold=True, color=LIRIE_GREEN)
        _compact_paragraph(p, space_before_pt=2, space_after_pt=2)
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(value)
            _apply_lirie_run(run, size_pt=BODY_SIZE_PT)
            _compact_paragraph(p, space_before_pt=2, space_after_pt=2)
    _add_para(doc, "", space_after_pt=4)


def _add_logo_header(doc: DocxDocument) -> bool:
    """Insère le logo LIRIE centré en tête de page 1. Retourne True si ajouté."""
    logo_path = _resolve_lirie_logo_path()
    if logo_path is None:
        return False
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _compact_paragraph(p, space_before_pt=0, space_after_pt=8, keep_with_next=True)
    run = p.add_run()
    # ~4.2 cm de large — visible sans saturer la page de couverture
    run.add_picture(str(logo_path), width=Cm(4.2))
    return True


def _section_page_break(doc: DocxDocument, *, label: str) -> None:
    """Saut de page contrôlé (pagination contractuelle recommandée)."""
    # label réservé au debug / lecture du code (ex. « page 3 — art. 5-6 »)
    _ = label
    doc.add_page_break()


def _fmt_days(value: Any, default: int) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return default


def _fmt_days_prose(
    value: int,
    *,
    calendar: bool = False,
    working: bool = False,
) -> str:
    """« trente (30) jours calendaires » — style rédactionnel juridique."""
    word = _num_to_word_fr(value)
    unit = "jour" if abs(value) <= 1 else "jours"
    suffix = ""
    if calendar:
        suffix = " calendaires"
    elif working:
        suffix = " ouvrables"
    return f"{word} ({value}) {unit}{suffix}"


def _penalty_text(commercial: dict[str, Any]) -> str:
    """Formule de la peine conventionnelle (hors commissions, dues par ailleurs).

    Retourne un complément de phrase à faire précéder de « au » (élision de
    « à » + « le ») par l'appelant, p. ex. : « ... correspondant au
    {_penalty_text(...)} », afin d'éviter toute construction fautive
    (« à le », « de les »).
    """
    penalty = commercial.get("penalty") or {}
    multiplier = penalty.get("multiplier", PENALTY_MULTIPLIER)
    minimum = penalty.get("minimum", PENALTY_MINIMUM_CHF)
    currency = penalty.get("currency", PENALTY_CURRENCY)
    minimum_fmt = f"{int(minimum):,}".replace(",", "'")
    return (
        f"montant le plus élevé entre deux fois ({multiplier}×) le montant "
        f"des commissions éludées et {currency} {minimum_fmt}.–"
    )


def build_partner_agreement_docx_bytes(
    *,
    reference: str,
    parties: dict[str, Any],
    commercial: dict[str, Any],
    agreement_effective_from: str,
) -> bytes:
    """Compose le contrat-cadre (A) + annexes B/C + signatures."""
    ensure_base_template()
    doc = Document(str(template_path()))

    # Nettoyer le corps du modèle de base
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)

    style = doc.styles["Normal"]
    style.font.name = LIRIE_FONT
    style.font.size = Pt(BODY_SIZE_PT)
    style.paragraph_format.space_after = Pt(SPACE_BODY_AFTER_PT)
    style.paragraph_format.line_spacing = LINE_SPACING
    _configure_lirie_styles(doc)

    _configure_section_and_footer(doc, reference=reference)

    operator = parties.get("operator") or {}
    partner = parties.get("partner") or {}
    signatory_authority = parties.get("signatory_authority_verification") or {}

    mode = commercial.get("subscription_pricing_mode") or "volume"
    free_months = commercial.get("free_license_max_months")
    commission_rate = commercial.get("commission_rate")
    cancel_policy = commercial.get("commission_cancellation_policy") or "exclude"
    payment_days = _fmt_days(commercial.get("payment_terms_days"), 30)
    dispute_days = _fmt_days(commercial.get("statement_dispute_days"), 10)
    commission_enabled = bool(commercial.get("lirie_commission_enabled", True))
    own_enabled = bool(commercial.get("own_portfolio_billing_enabled", True))

    # ——— Page 1 : logo + titre + identification des Parties ———
    _add_logo_header(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(
        "CONTRAT CADRE DE PARTENARIAT\n& LICENCE D'UTILISATION DE LA PLATEFORME LIRIE"
    )
    _apply_lirie_run(r, size_pt=13, bold=True, color=LIRIE_GREEN)
    _compact_paragraph(title, space_before_pt=0, space_after_pt=8, keep_with_next=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Plateforme digitale LIRIE  ·  www.lirie.ch")
    _apply_lirie_run(sr, size_pt=BODY_SIZE_PT, color=LIRIE_GREEN)
    _compact_paragraph(sub, space_before_pt=0, space_after_pt=10, keep_with_next=True)

    ref_p = doc.add_paragraph()
    ref_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = ref_p.add_run(f"Réf. : {reference}")
    _apply_lirie_run(rr, size_pt=BODY_SIZE_PT, bold=True, color=LIRIE_GREEN)
    _compact_paragraph(ref_p, space_before_pt=0, space_after_pt=6, keep_with_next=True)

    effective_date_fr = _fmt_effective_date_fr(agreement_effective_from)
    _add_para(
        doc,
        f"Date d'effet commerciale : {effective_date_fr}",
        keep_with_next=True,
        space_after_pt=4,
    )
    _add_para(
        doc,
        "Les conditions financières de la Partie B prennent effet à la date "
        "d'effet commerciale indiquée ci-dessus et peuvent s'appliquer aux "
        "opérations identifiables réalisées depuis cette date. Les autres "
        "dispositions contractuelles prennent effet à la date de la dernière "
        "signature des Parties, sauf lorsqu'une disposition prévoit "
        "expressément une autre date. La Partie C s'applique également aux "
        "données traitées dans le cadre de la relation depuis la date "
        "d'effet commerciale.",
        keep_with_next=True,
        space_after_pt=4,
    )
    _add_para(
        doc,
        f"Version modèle : {TEMPLATE_VERSION}",
        keep_with_next=True,
        space_after_pt=10,
    )

    _add_heading(doc, "ENTRE", level=1, space_before_pt=6, space_after_pt=4)
    op_lines = _party_block(role="L'Exploitant", party=operator, is_operator=True)
    for i, line in enumerate(op_lines):
        _add_para(
            doc,
            line,
            bold=(i == 0),
            keep_with_next=True,
            space_after_pt=SPACE_PARTY_LINE_AFTER_PT,
        )
    _add_para(
        doc,
        "Ci-après désigné : « l'Exploitant »",
        bold=True,
        keep_with_next=True,
        space_after_pt=10,
    )

    _add_heading(doc, "ET", level=1, space_before_pt=6, space_after_pt=4)
    pa_lines = _party_block(role="Le Partenaire", party=partner, is_operator=False)
    for i, line in enumerate(pa_lines):
        _add_para(
            doc,
            line,
            bold=(i == 0),
            keep_with_next=True,
            space_after_pt=SPACE_PARTY_LINE_AFTER_PT,
        )
    _add_para(
        doc,
        "Ci-après désigné : « le Partenaire »",
        bold=True,
        keep_with_next=True,
        space_after_pt=6,
    )
    _add_para(
        doc,
        "L'Exploitant et le Partenaire étant ensemble désignés « les Parties ».",
        keep_with_next=False,
        space_after_pt=8,
    )

    # ——— PARTIE A — CONTRAT-CADRE ———
    _section_page_break(doc, label="page 2 — Partie A, art. 1 à 5")
    _add_heading(doc, "PARTIE A — CONTRAT-CADRE", level=1)
    _add_para(
        doc,
        "La présente Partie A régit le cadre juridique de la relation entre les "
        "Parties. Les conditions commerciales chiffrées figurent en Partie B "
        "(Annexe financière) et les règles de protection des données en Partie C "
        "(Annexe de protection des données) ; ces trois parties forment un ensemble "
        "contractuel indivisible.",
    )

    _add_heading(doc, "ARTICLE 1 – DÉFINITIONS", level=1)
    _add_bullets(
        doc,
        [
            "« Plateforme » : l'application et les services numériques LIRIE "
            "permettant l'enregistrement, la mise à disposition et la transmission "
            "de demandes de transport, la gestion des comptes, la facturation et le "
            "suivi des courses.",
            "« Partenaire » : le prestataire de transport professionnel partie au "
            "présent contrat, identifié en page 1.",
            "« Course du portefeuille propre » : toute course enregistrée par le "
            "Partenaire dans la Plateforme pour un client ou une relation "
            "commerciale qu'il apporte et gère lui-même, indépendamment de toute "
            "mise en relation par LIRIE (origine technique OWN_PORTFOLIO).",
            "« Course transmise par LIRIE » : toute demande de transport créée par "
            "une institution, un client ou un autre utilisateur, transmise au "
            "Partenaire par l'intermédiaire de la Plateforme et expressément "
            "acceptée par celui-ci (origine technique LIRIE_MARKETPLACE). Son "
            "exécution, son annulation ou sa facturation ultérieure n'altère pas "
            "son origine ; ces événements déterminent uniquement son traitement "
            "financier conformément à la Partie B.",
            "« Relevé » : le document périodique récapitulant les montants dus "
            "entre les Parties, conformément à la Partie B.",
            "« Passager » : la personne physique effectivement transportée lors "
            "d'une course, qu'elle soit ou non à l'origine de la demande.",
            "« Demandeur » : la personne ou l'entité qui crée ou enregistre une "
            "demande de transport dans la Plateforme, qu'elle en soit ou non le "
            "Client contractuel.",
            "« Client contractuel » : la personne physique ou morale qui conclut "
            "le contrat de transport avec le Partenaire, conformément à "
            "l'article 6.",
            "« Payeur désigné » : la personne ou l'entité indiquée dans la "
            "Plateforme comme devant régler la course ; cette seule indication "
            "ne vaut pas acceptation, par cette personne ou entité, de payer la "
            "course.",
            "« Destinataire de la facture » : la personne ou l'entité à "
            "laquelle la facture de la course est adressée, qui peut différer "
            "du Client contractuel ou du Payeur désigné.",
            "« Utilisateur » : toute personne physique accédant à la "
            "Plateforme, quel que soit son rôle (Partenaire, collaborateur du "
            "Partenaire, Passager, Demandeur, institution ou autre).",
        ],
    )

    _add_heading(doc, "ARTICLE 2 – OBJET ET INDÉPENDANCE DES PARTIES", level=1)
    _add_para(
        doc,
        "Le présent contrat a pour objet l'octroi au Partenaire d'un droit "
        "d'utilisation professionnel de la Plateforme, l'organisation de la "
        "transmission des demandes de transport, la fixation des conditions "
        "financières applicables (Partie B) et la répartition des "
        "responsabilités respectives des Parties.",
    )
    _add_para(
        doc,
        "Les Parties agissent en qualité d'entités indépendantes. Le présent "
        "contrat ne constitue ni un contrat de travail, ni une société simple au "
        "sens des art. 530 ss CO, ni un contrat d'agence ou de représentation au "
        "sens des art. 418a ss CO, ni un mandat exclusif. Chaque Partie assume "
        "seule ses charges sociales, fiscales et d'assurance.",
    )

    _add_heading(doc, "ARTICLE 3 – STATUT DE LIRIE", level=1)
    _add_para(
        doc,
        "LIRIE intervient comme fournisseur de la Plateforme et, pour les Courses "
        "transmises par LIRIE, comme intermédiaire technique de mise en relation. "
        "Elle n'est ni transporteur, ni employeur des chauffeurs, ni partie au "
        "contrat de transport conclu entre le Partenaire et le Client "
        "contractuel.",
    )
    _add_para(
        doc,
        "Pour l'enregistrement, la mise à disposition et la transmission des "
        "demandes, LIRIE est tenue d'une obligation de moyens : elle met en "
        "œuvre les moyens techniques raisonnables pour assurer le bon "
        "fonctionnement de la Plateforme, sans garantir la réception effective "
        "d'une demande par le dispositif du Partenaire, ni son acceptation par "
        "le Partenaire, ni la bonne exécution matérielle du transport, qui "
        "relève exclusivement du Partenaire. Il appartient au Partenaire de "
        "vérifier, dans son tableau de bord LIRIE, l'état et le contenu de "
        "chaque demande avant toute exécution.",
    )
    _add_para(
        doc,
        "Les journaux techniques (logs) de la Plateforme constituent une "
        "présomption simple du déroulement des opérations système "
        "enregistrées (création, mise à disposition, envoi d'une "
        "notification, actions sur le compte), susceptible d'être renversée "
        "par preuve contraire ; ils ne constituent pas une garantie de la "
        "prise de connaissance effective de ces éléments par leur "
        "destinataire.",
    )

    _add_heading(doc, "ARTICLE 4 – LICENCE, COMPTES ET SÉCURITÉ", level=1)
    _add_para(
        doc,
        "LIRIE accorde au Partenaire une licence non exclusive, non cessible, "
        "strictement professionnelle et limitée à la durée du présent contrat.",
    )
    _add_para(
        doc,
        "Les comptes d'accès à la Plateforme sont nominatifs. Le Partenaire "
        "interdit tout partage d'identifiants entre utilisateurs et répond des "
        "actions effectuées depuis ses comptes. Il signale à LIRIE, sans délai, "
        "toute compromission ou usage non autorisé constaté ; LIRIE peut alors "
        "suspendre l'accès concerné à titre conservatoire.",
    )

    _add_heading(
        doc,
        "ARTICLE 5 – PORTEFEUILLE PROPRE, COURSES TRANSMISES ET "
        "RÉMUNÉRATION DE FACILITATION",
        level=1,
    )
    _add_para(
        doc,
        "Le portefeuille propre et les Courses transmises par LIRIE (définis à "
        "l'article 1) font l'objet d'un traitement financier distinct, précisé "
        "en Partie B.",
    )
    _add_para(
        doc,
        "Pour toute Course transmise par LIRIE, le Partenaire doit à LIRIE "
        "une rémunération de facilitation (« commission ») calculée selon le "
        "résultat financier définitif de la course et conformément aux "
        "conditions de la Partie B.",
    )
    _add_para(
        doc,
        "L'existence d'une relation commerciale préexistante entre le Partenaire "
        "et un client n'exclut pas la commission dès lors qu'une demande "
        "concernant ce client est transmise par la Plateforme et y est "
        "enregistrée ; le contournement de cette règle est sanctionné à "
        "l'article 8.",
    )

    _section_page_break(doc, label="page 3 — Partie A, art. 6 à 8")
    _add_heading(doc, "ARTICLE 6 – FORMATION DU CONTRAT DE TRANSPORT", level=1)
    _add_para(
        doc,
        "Le contrat de transport se forme entre le Partenaire et le Client "
        "contractuel au moment de l'acceptation définitive de la demande par "
        "le Partenaire dans la Plateforme — et non à l'exécution de la "
        "course, ni à sa facturation.",
    )
    _add_para(
        doc,
        "La seule désignation d'un Payeur ou d'un Destinataire de facture ne "
        "crée aucune obligation de paiement. Une telle obligation ne peut "
        "résulter que d'un engagement séparé, exprès ou autrement démontrable, "
        "de la personne ou de l'entité concernée.",
    )
    _add_para(
        doc,
        "Avant de procéder à l'acceptation définitive, le Partenaire vérifie les "
        "informations essentielles affichées (adresses, horaire, nature de la "
        "prestation, contraintes éventuelles). En cas d'incertitude sur "
        "l'identité du Client contractuel ou sur le fondement de l'obligation "
        "du Payeur, le Partenaire demande une clarification avant d'accepter ; "
        "il lui appartient de refuser ou de faire préciser toute demande "
        "incomplète ou ambiguë.",
    )
    _add_para(
        doc,
        "Constitue notamment une modification substantielle de la demande, "
        "postérieure à l'acceptation définitive : le changement du lieu de "
        "prise en charge ou de destination, de la date ou de l'heure, "
        "l'ajout d'un trajet retour, le changement de type de véhicule, la "
        "nécessité d'un accès en fauteuil roulant ou d'oxygène, la présence "
        "d'un accompagnant, le changement du nombre de passagers, le "
        "changement de Client contractuel ou de Payeur désigné, ou une "
        "modification tarifaire.",
    )
    _add_para(
        doc,
        "Tant que le Partenaire n'a pas accepté une telle modification, "
        "celle-ci ne lui est pas opposable ; la demande initiale demeure "
        "régie par les règles d'annulation et de modification applicables.",
    )

    _add_heading(doc, "ARTICLE 7 – OBLIGATIONS DU PARTENAIRE", level=1)
    _add_para(doc, "Le Partenaire s'engage notamment à :")
    _add_bullets(
        doc,
        [
            "exécuter personnellement, ou par l'intermédiaire de son personnel "
            "qualifié, les transports acceptés, dans le respect des lois "
            "suisses applicables au transport de personnes ;",
            "maintenir en vigueur les autorisations administratives et "
            "assurances requises (article 9) ;",
            "renseigner dans la Plateforme, de manière exacte et complète, le "
            "montant HT définitif facturé et les informations nécessaires à "
            "l'établissement du Relevé, conserver les justificatifs "
            "correspondants, déclarer toute correction, note de crédit ou "
            "annulation ultérieure, et s'interdire de réduire artificiellement "
            "un montant afin de diminuer la commission due ;",
            "s'acquitter des sommes dues à LIRIE aux échéances fixées en Partie B ;",
            "informer LIRIE sans délai de toute circonstance affectant "
            "significativement sa capacité à exécuter les transports acceptés.",
        ],
    )

    _add_heading(doc, "ARTICLE 8 – NON-CONTOURNEMENT", level=1)
    _add_para(
        doc,
        "Pendant la durée du contrat et pendant six (6) mois après sa "
        "résiliation, le Partenaire s'interdit de solliciter activement un "
        "client qui lui a été présenté pour la première fois par "
        "l'intermédiaire de la Plateforme, dans le but de soustraire à LIRIE des "
        "courses qui auraient dû y être enregistrées. L'interdiction de "
        "non-contournement ne prive pas le Partenaire de ses relations "
        "commerciales préexistantes indépendantes de LIRIE, ni des appels "
        "d'offres publics ou ouverts.",
    )
    _add_para(
        doc,
        "Toutefois, toute demande reçue directement qui constitue la suite, "
        "le renouvellement ou la répétition d'une demande initialement "
        "transmise ou gérée par LIRIE doit être enregistrée dans la "
        "Plateforme pendant la durée du contrat et la période de protection "
        "applicable. Toute demande spontanée émanant d'un client présenté "
        "pour la première fois par LIRIE doit également être enregistrée "
        "afin d'en préserver la traçabilité.",
    )
    _add_para(
        doc,
        "Toute violation intentionnelle du présent article, ou sciemment "
        "dissimulée, constitue un manquement grave. Le Partenaire demeure "
        "tenu de payer l'intégralité des commissions éludées. Il doit en "
        "outre une peine conventionnelle correspondant "
        f"au {_penalty_text(commercial)}. Une seule peine conventionnelle "
        "est appliquée à un même ensemble de faits, indépendamment du "
        "nombre de courses concernées. Tout dommage supplémentaire ne peut "
        "être réclamé que dans la mesure où il n'est pas déjà couvert par "
        "cette peine. Le pouvoir du juge de réduire une peine conventionnelle "
        "excessive demeure réservé.",
    )
    _add_para(
        doc,
        "Les sanctions et l'interdiction de non-contournement prévues au "
        "présent article s'appliquent à compter de la date de la dernière "
        "signature du présent contrat par les Parties, indépendamment de la "
        "date d'entrée en vigueur commerciale retenue en Partie B.",
    )

    _add_heading(doc, "ARTICLE 9 – ASSURANCES", level=1)
    _add_para(
        doc,
        "Le Partenaire dispose en permanence d'une assurance responsabilité "
        "civile professionnelle, d'assurances véhicules adaptées, d'une "
        "couverture passagers et des autorisations administratives "
        "nécessaires à son activité. Il fournit à première demande, et au "
        "minimum une fois par année, une attestation d'assurance à jour, et "
        "informe immédiatement LIRIE de toute suspension, réduction ou "
        "résiliation de couverture.",
    )

    _section_page_break(doc, label="page 4 — Partie A, art. 10 à 12")
    _add_heading(doc, "ARTICLE 10 – RESPONSABILITÉ", level=1)
    _add_para(
        doc,
        "Dans la mesure permise par la loi, et sauf dol ou faute grave, LIRIE "
        "n'assume aucune responsabilité pour faute légère s'agissant : de "
        "l'exécution du transport, des chauffeurs et véhicules du Partenaire, "
        "de la fixation ou de l'encaissement des prix, des moyens de paiement "
        "utilisés, de l'exactitude des données transmises par des tiers "
        "(clients, institutions), de la configuration choisie par le Partenaire "
        "dans la Plateforme, ainsi que du fait de tiers échappant à son "
        "contrôle raisonnable.",
    )
    _add_para(
        doc,
        "Lorsque la responsabilité de LIRIE est engagée, seuls les dommages "
        "directs, prouvés et raisonnablement prévisibles lors de la conclusion "
        "du présent contrat sont indemnisables. Sont exclus tout gain manqué, "
        "toute perte de chiffre d'affaires, de clientèle ou d'opportunité "
        "commerciale, ainsi que tout dommage indirect ou toute atteinte à la "
        "réputation.",
    )
    _add_para(
        doc,
        "Sauf dol ou faute grave — non susceptibles de limitation —, la "
        "responsabilité totale et cumulée de LIRIE au titre du présent contrat "
        "est plafonnée, à titre subsidiaire, au montant total hors taxes des "
        "commissions et abonnements effectivement payés par le Partenaire au "
        "cours des douze (12) mois précédant le fait générateur.",
    )

    _add_heading(doc, "ARTICLE 11 – INDEMNISATION PAR LE PARTENAIRE", level=1)
    _add_para(
        doc,
        "Le Partenaire demeure seul responsable, envers les clients, tiers et "
        "autorités, de l'exécution des transports, de son personnel, de ses "
        "véhicules et du respect des lois applicables. Il garantit LIRIE "
        "contre toute réclamation de tiers résultant d'un manquement du "
        "Partenaire à ses obligations légales ou contractuelles et la relève "
        "des conséquences financières raisonnablement démontrées en "
        "découlant, frais de défense raisonnables inclus, sous réserve de son "
        "droit d'en contester le bien-fondé.",
    )
    _add_para(
        doc,
        "LIRIE notifie le Partenaire dans un délai raisonnable de toute "
        "réclamation de tiers susceptible de donner lieu à la garantie "
        "prévue au présent article. Le Partenaire peut participer à sa "
        "défense, à ses frais, en coordination avec LIRIE.",
    )
    _add_para(
        doc,
        "Aucune transaction ou reconnaissance de responsabilité affectant le "
        "Partenaire n'est conclue par LIRIE sans l'accord écrit préalable de "
        "celui-ci, accord qui ne peut être refusé sans motif raisonnable.",
    )
    _add_para(
        doc,
        "La présente garantie ne couvre pas la part de responsabilité "
        "imputable à une faute propre de LIRIE. En cas de responsabilité "
        "concurrente des Parties, la charge en est répartie "
        "proportionnellement à la part de faute respective de chacune.",
    )

    _add_heading(
        doc, "ARTICLE 12 – CONFIDENTIALITÉ ET PROPRIÉTÉ INTELLECTUELLE", level=1
    )
    _add_para(
        doc,
        "Chaque Partie conserve confidentielles les informations non "
        "publiques de l'autre Partie dont elle a connaissance dans le cadre "
        "du présent contrat. Cette obligation demeure en vigueur pendant "
        "toute la durée du contrat et, pour les secrets d'affaires, tant que "
        "l'information conserve son caractère secret ; le traitement des "
        "données personnelles est régi par la Partie C. Sont exclues les "
        "informations tombées dans le domaine public sans violation du "
        "présent article, ainsi que celles déjà connues licitement ou "
        "développées indépendamment par la Partie destinataire.",
    )
    _add_para(
        doc,
        "Chaque Partie répond du respect de la présente obligation par ses "
        "collaborateurs, employés et mandataires auxquels elle donnerait "
        "accès à ces informations. Une divulgation imposée par la loi, une "
        "autorité ou une décision judiciaire est autorisée, dans la limite "
        "strictement nécessaire et, dans la mesure du possible, après "
        "information préalable de l'autre Partie.",
    )
    _add_para(
        doc,
        "À l'issue du contrat, chaque Partie restitue ou détruit, sur "
        "demande écrite de l'autre Partie, les informations confidentielles "
        "reçues, sous réserve des obligations légales de conservation et des "
        "copies nécessaires à la preuve de ses droits.",
    )
    _add_para(
        doc,
        "La Plateforme (code, architecture, algorithmes, interfaces, marque, "
        "documentation) demeure la propriété exclusive de LIRIE ; toute "
        "amélioration ou tout développement réalisé par LIRIE à partir de "
        "l'utilisation de la Plateforme par le Partenaire lui demeure "
        "acquis. Ni le Partenaire ni LIRIE ne peut utiliser le nom, le logo "
        "ou la marque de l'autre Partie sans son accord écrit préalable. "
        "LIRIE peut produire des statistiques à partir de données "
        "irréversiblement anonymisées ou agrégées de manière à ne permettre "
        "raisonnablement ni l'identification ni la réidentification d'une "
        "personne, destinées à améliorer ses services.",
    )

    _add_heading(doc, "ARTICLE 13 – DISPONIBILITÉ ET SUPPORT", level=1)
    _add_para(
        doc,
        "LIRIE s'engage, par obligation de moyens, à maintenir la Plateforme "
        "accessible et fonctionnelle, sans garantie de disponibilité absolue "
        "(pas de SLA contractuel). Les opérations de maintenance sont annoncées "
        "lorsque cela est raisonnablement possible.",
    )
    _add_para(
        doc,
        "L'assistance ordinaire nécessaire à l'utilisation des "
        "fonctionnalités existantes de la Plateforme est incluse, dans des "
        "limites raisonnables. Les modalités spécifiques (support "
        "facturable, configurations particulières, formations, "
        "développements sur demande) sont précisées, le cas échéant, en "
        "Partie B. Tout développement spécifique fait l'objet d'un devis "
        "préalable accepté par écrit par le Partenaire.",
    )

    _section_page_break(doc, label="page 5 — Partie A, art. 14-16")
    _add_heading(doc, "ARTICLE 14 – DURÉE ET RÉSILIATION", level=1)
    _add_para(
        doc,
        "Le contrat est conclu pour une durée indéterminée dès son entrée en "
        "vigueur. Chaque Partie peut résilier le contrat à tout moment "
        "moyennant un préavis écrit de trente (30) jours calendaires, "
        "adressé à l'autre Partie.",
    )
    _add_para(
        doc,
        "En cas de défaut de paiement, LIRIE applique la procédure progressive "
        "de rappel puis de suspension définie en Partie B avant toute "
        "résiliation fondée sur ce motif.",
    )
    _add_para(
        doc,
        "Une résiliation immédiate, sans respecter le préavis ordinaire, est "
        "possible en cas de faute grave — notamment contournement (article 8), "
        "fraude, absence d'assurance obligatoire ou atteinte grave à la "
        "sécurité des données — après mise en demeure restée sans effet, sauf "
        "urgence ou fraude avérée.",
    )
    _add_para(
        doc,
        "Quel qu'en soit le motif, à l'issue du contrat le Partenaire conserve "
        "un accès en lecture seule à ses données pendant trente (30) jours afin "
        "d'en exporter le contenu. Les commissions et abonnements déjà dus "
        "restent exigibles ; les clauses de confidentialité, de propriété "
        "intellectuelle, de protection des données (Partie C) et de "
        "non-contournement (article 8) survivent selon leurs termes propres.",
    )

    _add_heading(
        doc,
        "ARTICLE 15 – FORCE MAJEURE, CESSION, DROIT APPLICABLE ET HIÉRARCHIE",
        level=1,
    )
    _add_para(
        doc,
        "Aucune Partie n'est responsable en cas d'événement imprévisible et "
        "indépendant de sa volonté. La Partie empêchée informe l'autre sans "
        "délai et prend les mesures raisonnables pour limiter les effets. Si "
        "l'empêchement se prolonge au-delà de soixante (60) jours, chaque "
        "Partie peut résilier le contrat sans indemnité.",
    )
    _add_para(
        doc,
        "Le Partenaire accepte que le présent contrat puisse être cédé à toute "
        "société ultérieurement constituée pour exploiter LIRIE, à une société "
        "du groupe, ou à une entité issue d'une fusion ou restructuration. La "
        "cession est notifiée au Partenaire. Le cessionnaire reprend les "
        "droits et obligations futurs ainsi que, dans la mesure permise par "
        "le droit applicable et prévue dans l'acte de reprise notifié au "
        "Partenaire, les droits et obligations déjà nés. La libération de "
        "l'Exploitant initial pour les obligations antérieures n'intervient "
        "que dans la mesure où elle est juridiquement valable et acceptée "
        "par le Partenaire.",
    )
    _add_para(
        doc,
        "Le présent contrat est soumis au droit matériel suisse, à "
        "l'exclusion de ses règles de conflit de lois. Sous réserve des fors "
        "impératifs, les tribunaux ordinaires du canton de Genève sont "
        "exclusivement compétents. La version française du contrat prévaut sur "
        "toute traduction.",
    )
    _add_para(
        doc,
        "Tout avenant signé postérieurement prévaut sur les dispositions "
        "qu'il modifie expressément. La Partie B prévaut sur la Partie A "
        "uniquement pour les questions financières. La Partie C prévaut sur "
        "les Parties A et B pour toute question relative à la protection "
        "des données. La Partie A s'applique pour le surplus.",
    )

    _add_heading(doc, "ARTICLE 16 – DISPOSITIONS FINALES", level=1)
    _add_bullets(
        doc,
        [
            "Invalidité partielle : si une clause du présent contrat est "
            "jugée invalide ou inapplicable, les autres clauses restent "
            "pleinement en vigueur ; les Parties s'efforcent de remplacer la "
            "clause invalide par une disposition valable de portée "
            "économique équivalente.",
            "Absence de renonciation : le fait pour une Partie de ne pas se "
            "prévaloir d'un manquement ou d'un droit ne constitue pas une "
            "renonciation à s'en prévaloir ultérieurement.",
            "Intégralité de l'accord : le présent contrat (Parties A, B et "
            "C) constitue l'intégralité de l'accord entre les Parties sur "
            "son objet et remplace tout accord ou échange antérieur portant "
            "sur le même objet.",
            "Modifications : toute modification du présent contrat doit "
            "être convenue par écrit entre les Parties, sous réserve des "
            "ajustements prévus par la Partie B lorsqu'ils sont expressément "
            "autorisés par le contrat.",
            "Notifications : les communications contractuelles sont "
            "valablement adressées aux coordonnées de contact indiquées par "
            "les Parties ; chaque Partie informe l'autre sans délai de toute "
            "modification de ses coordonnées.",
            "Signature : le présent contrat peut être signé de manière "
            "manuscrite ou électronique, en plusieurs exemplaires ou copies "
            "ayant chacun valeur d'original.",
            "Survie des clauses : les clauses relatives à la "
            "confidentialité et à la propriété intellectuelle (article 12), "
            "à la protection des données (Partie C), au non-contournement "
            "(article 8), à la responsabilité (article 10) et à "
            "l'indemnisation (article 11) survivent à la fin du contrat, "
            "quelle qu'en soit la cause.",
        ],
    )

    # ——— PARTIE B — ANNEXE FINANCIÈRE ———
    _section_page_break(doc, label="page 6 — Partie B, annexe financière")
    _add_heading(doc, "PARTIE B — ANNEXE FINANCIÈRE", level=1)
    _add_para(
        doc,
        "La présente annexe fixe les conditions financières applicables au "
        "Partenaire. Elle fait partie intégrante du contrat et prévaut sur "
        "la Partie A pour les questions financières, conformément à "
        "l'article 15.",
    )

    _add_heading(doc, "B.1 Produits activés", level=2)
    _add_bullets(
        doc,
        [
            f"Abonnement portefeuille propre : "
            f"{'activé' if own_enabled else 'non activé'} ;",
            f"Commission sur les Courses transmises par LIRIE : "
            f"{'activée' if commission_enabled else 'non activée'} ;",
            f"Support facturable : "
            f"{'activé' if bool(commercial.get('support_enabled')) else 'non activé'}.",
        ],
    )

    _add_heading(doc, "B.2 Commission sur les Courses transmises", level=2)
    if commission_enabled:
        due_if_unpaid = bool(commercial.get("commission_due_if_customer_unpaid", True))
        _add_para(
            doc,
            f"Taux applicable : {_pct(commission_rate)} du montant HT définitif "
            "facturé au titre de la prestation de transport, y compris les "
            "suppléments directement liés à la course, après déduction des "
            "remises, rabais, remboursements et notes de crédit. Sont exclus "
            "les pourboires, débours remboursés au prix coûtant, taxes "
            "publiques et montants de TVA.",
        )
        _add_para(
            doc,
            f"Politique d'annulation : "
            f"{_CANCEL_LABELS.get(cancel_policy, cancel_policy)}.",
        )
        _add_para(
            doc,
            "Le Partenaire renseigne le montant HT définitif dans les cinq "
            "jours ouvrables suivant l'émission de sa facture au client. "
            "Toute note de crédit, annulation ou correction ultérieure est "
            "régularisée sur le Relevé suivant. En présence d'une "
            "divergence concrète ou d'un indice sérieux, LIRIE peut demander "
            "les justificatifs strictement nécessaires à la vérification.",
        )
        if due_if_unpaid:
            _add_para(
                doc,
                "La commission reste due à LIRIE même si le client final ne "
                "règle pas, ou ne règle que partiellement, le montant facturé "
                "par le Partenaire ; le risque d'encaissement client relève de "
                "la relation entre le Partenaire et son client.",
            )
    else:
        _add_para(
            doc,
            "Aucune commission n'est due tant que ce produit n'est pas activé "
            "dans les conditions commerciales.",
        )

    _add_heading(doc, "B.3 Abonnement portefeuille propre", level=2)
    if own_enabled:
        mode_label = _MODE_LABELS.get(mode, mode)
        _add_para(doc, f"Mode tarifaire applicable : {mode_label}.")
        if mode == "free":
            months_txt = str(free_months) if free_months else "60"
            _add_para(
                doc,
                f"La licence du portefeuille propre est gratuite pendant "
                f"{months_txt} mois calendaires à compter du "
                f"{effective_date_fr}, pour autant que le présent contrat "
                "demeure en vigueur. Aucun abonnement ni frais fixe "
                "d'utilisation n'est dû pour cette licence durant cette "
                "période ; la gratuité ne concerne pas la commission sur les "
                "Courses transmises par LIRIE.",
            )
            _add_para(
                doc,
                f"À l'expiration des {months_txt} mois de gratuité, "
                "d'éventuelles conditions tarifaires devront faire l'objet "
                "d'un nouvel avenant à la présente Annexe financière, signé "
                "par les deux Parties. Aucun abonnement ne sera appliqué "
                "automatiquement en l'absence d'un tel avenant.",
            )
        elif mode == "fixed":
            amount = commercial.get("custom_subscription_amount")
            amount_txt = (
                f"CHF {_fmt_chf_amount(amount)}"
                if amount
                else "le montant fixe convenu"
            )
            _add_para(
                doc,
                f"L'abonnement portefeuille propre est facturé à hauteur de "
                f"{amount_txt} par mois, hors taxes sauf mention contraire "
                "(cf. B.7 TVA).",
            )
        else:
            _add_para(
                doc,
                "L'abonnement portefeuille propre est facturé selon le volume "
                "mensuel de courses, conformément à la grille de paliers "
                "ci-dessous.",
            )
            pricing = commercial.get("subscription_pricing") or {}
            tiers = pricing.get("tiers") or []
            if tiers:
                rows = []
                for tier in tiers:
                    vmin = tier.get("volume_min")
                    vmax = tier.get("volume_max")
                    range_txt = (
                        f"Dès {vmin} courses"
                        if vmax is None
                        else f"{vmin} à {vmax} courses"
                    )
                    price = tier.get("price_monthly")
                    price_txt = (
                        f"CHF {_fmt_chf_amount(price)}/mois"
                        if price is not None
                        else "—"
                    )
                    rows.append([range_txt, tier.get("label") or "—", price_txt])
                _add_simple_table(
                    doc,
                    headers=["Palier de volume", "Libellé", "Prix mensuel"],
                    rows=rows,
                )
    else:
        _add_para(
            doc,
            "La facturation de l'abonnement portefeuille propre n'est pas "
            "activée pour ce Partenaire.",
        )

    _add_heading(doc, "B.4 Support", level=2)
    if bool(commercial.get("support_enabled")):
        hourly = commercial.get("support_hourly_rate_default")
        rate_txt = _fmt_chf_amount(hourly) if hourly else "tarif convenu"
        _add_para(
            doc,
            "L'assistance ordinaire nécessaire à l'utilisation des "
            "fonctionnalités existantes est incluse, dans des limites "
            "raisonnables. Les prestations spécifiques, configurations "
            "particulières, formations supplémentaires, interventions "
            "imputables à une mauvaise utilisation ou développements sur "
            f"demande sont facturés à CHF {rate_txt} par heure, hors TVA. Le "
            "temps consacré est documenté et communiqué au Partenaire. Il "
            "est réputé accepté s'il n'est pas contesté de manière motivée "
            "dans les cinq jours ouvrables suivant sa notification.",
        )
    else:
        _add_para(
            doc,
            "L'assistance ordinaire nécessaire à l'utilisation des "
            "fonctionnalités existantes est incluse, dans des limites "
            "raisonnables. Aucun produit de support facturable n'est activé "
            "pour ce Partenaire.",
        )

    _add_heading(doc, "B.5 Modalités de paiement et contestation", level=2)
    _add_para(
        doc,
        f"Relevé mensuel ; facturation par LIRIE ; paiement sous "
        f"{_fmt_days_prose(payment_days)} à compter de la date d'émission "
        "de la facture ; intérêt moratoire de 5 % l'an (art. 104 CO) en cas "
        "de retard. Le Partenaire dispose de "
        f"{_fmt_days_prose(dispute_days, working=True)} à compter de la "
        "notification de la mise à disposition du relevé dans LIRIE pour le "
        "contester de manière motivée ; à défaut, le relevé est réputé accepté "
        "sous réserve d'erreur manifeste. La contestation d'une partie du "
        "relevé ne suspend pas le paiement de la partie non contestée. Le "
        "Partenaire est en demeure automatiquement à l'échéance indiquée sur "
        "la facture, sans rappel préalable.",
    )

    auto_dunning = bool(commercial.get("automated_dunning_enabled", True))
    _add_heading(doc, "B.6 Procédure en cas de défaut de paiement", level=2)
    if not auto_dunning:
        _add_para(
            doc,
            "Les mesures automatisées de rappel et de suspension ne sont pas "
            "activées pour le présent contrat. LIRIE conserve ses droits de "
            "réclamer les montants échus, les intérêts moratoires, de "
            "suspendre les services après notification et d'engager les "
            "démarches de recouvrement prévues par le droit applicable.",
        )
    else:
        rem_delay = _fmt_days(commercial.get("reminder_delay_days_after_due"), 0)
        rem_grace = _fmt_days(commercial.get("reminder_grace_days"), 10)
        full_days = _fmt_days(commercial.get("full_suspend_days_after_due"), 30)
        full_count = _fmt_days(commercial.get("full_suspend_overdue_invoice_count"), 2)
        term_days = _fmt_days(commercial.get("termination_notice_days"), 10)
        block_offers = bool(commercial.get("partial_block_marketplace_offers", True))
        block_accept = bool(
            commercial.get("partial_block_marketplace_acceptance", True)
        )
        block_support = bool(commercial.get("partial_block_billable_support", True))
        block_config = bool(
            commercial.get("partial_block_billable_configuration", True)
        )
        _add_para(
            doc,
            f"LIRIE adresse au Partenaire une notification (délai de "
            f"{_fmt_days_prose(rem_delay)} après l'échéance) lui accordant "
            f"un délai supplémentaire de {_fmt_days_prose(rem_grace, calendar=True)} "
            "pour régler l'intégralité des montants échus ou contester la "
            "facture de manière motivée.",
        )
        bullets: list[str] = []
        if block_offers or block_accept:
            bullets.append(
                "suspendre la transmission et/ou l'acceptation de nouvelles "
                "Courses transmises par LIRIE ;"
            )
        if block_support or block_config:
            bullets.append(
                "suspendre l'accès aux fonctionnalités facturables et refuser "
                "les prestations de support ou de configuration facturables ;"
            )
        bullets.append("engager une procédure de recouvrement ou de poursuite.")
        _add_para(
            doc,
            "Si le paiement n'est pas intervenu à l'expiration de ce délai, "
            "LIRIE peut, après en avoir informé le Partenaire :",
        )
        _add_bullets(doc, bullets)
        _add_para(
            doc,
            "La suspension ne fait pas obstacle à l'exécution et au suivi des "
            "courses déjà engagées, ni à l'accès temporaire aux informations "
            "nécessaires à leur réalisation, à la sécurité des passagers, au "
            "paiement des factures ou à l'export des données du Partenaire.",
        )
        full_count_txt = _num_to_word_fr(full_count)
        _add_para(
            doc,
            f"Lorsque le retard dépasse {_fmt_days_prose(full_days)}, lorsque "
            f"{full_count_txt} factures échues et impayées demeurent "
            "ouvertes, ou lorsque le Partenaire ne respecte pas un accord "
            "de paiement "
            "écrit, LIRIE peut restreindre l'ensemble des opérations "
            "commerciales nouvelles et résilier le contrat pour faute grave, "
            f"après une dernière mise en demeure restée sans effet pendant "
            f"{_fmt_days_prose(term_days)}.",
        )
    _add_para(
        doc,
        "Les frais officiels de poursuite ainsi que tout dommage "
        "supplémentaire nécessaire et démontré résultant du retard peuvent "
        "être réclamés au Partenaire conformément au droit applicable. Aucun "
        "frais forfaitaire de recouvrement non justifié n'est automatiquement "
        "dû. LIRIE peut, sans y être obligée, accepter un échéancier écrit, "
        "sans que cela ne vaille remise de dette ni renonciation à ses autres "
        "droits.",
    )

    _add_heading(doc, "B.7 Pénalité de non-contournement (rappel)", level=2)
    _add_para(
        doc,
        "Conformément à l'article 8 de la Partie A, toute violation de "
        "l'interdiction de non-contournement oblige le Partenaire à payer "
        "l'intégralité des commissions éludées. Il doit en outre une peine "
        f"conventionnelle correspondant au {_penalty_text(commercial)}. Le "
        "pouvoir du juge de réduire une peine conventionnelle excessive "
        "demeure réservé.",
    )

    _add_heading(doc, "B.8 TVA et présentation des montants", level=2)
    if bool(commercial.get("amounts_are_tax_inclusive")):
        _add_para(
            doc,
            "Sauf mention contraire, les montants indiqués dans les relevés et "
            "factures s'entendent toutes taxes comprises (TTC).",
        )
    else:
        _add_para(
            doc,
            "Sauf mention contraire, les montants indiqués dans les relevés et "
            "factures s'entendent hors taxes (HT). La TVA est ajoutée, le cas "
            "échéant, au taux légal en vigueur à la date de facturation.",
        )

    special_conditions = (commercial.get("contract_special_conditions") or "").strip()
    if special_conditions:
        _add_heading(doc, "B.9 Conditions particulières", level=2)
        for line in special_conditions.splitlines():
            if line.strip():
                _add_para(doc, line.strip())

    _add_para(
        doc,
        f"Les présentes conditions commerciales s'appliquent à compter du "
        f"{effective_date_fr}.",
        space_before_pt=6,
    )

    # ——— PARTIE C — ANNEXE DE PROTECTION DES DONNÉES ———
    _section_page_break(doc, label="page 7 — Partie C, protection des données")
    _add_heading(doc, "PARTIE C — ANNEXE DE PROTECTION DES DONNÉES", level=1)
    _add_para(
        doc,
        "La présente annexe précise, dans le respect de la loi fédérale sur la "
        "protection des données (LPD), la répartition des rôles entre LIRIE et "
        "le Partenaire. Une institution ou un client à l'origine d'une demande "
        "n'est pas partie au présent contrat ; elle peut néanmoins agir en "
        "qualité de responsable du traitement pour ses propres finalités.",
    )
    compliance = commercial.get("compliance") or {}
    retention_ver = (
        compliance.get("retention_policy_version") or RETENTION_POLICY_VERSION
    )
    subprocessors_ver = compliance.get("subprocessors_version") or SUBPROCESSORS_VERSION
    penalty_ver = (
        compliance.get("penalty_calculation_version") or PENALTY_CALCULATION_VERSION
    )
    _add_para(
        doc,
        f"Versions de référence : politique de conservation ({retention_ver}) ; "
        f"liste des sous-traitants ({subprocessors_ver}) ; calcul de la "
        f"pénalité ({penalty_ver}).",
    )

    _add_heading(doc, "C.1 Répartition des rôles", level=2)
    _add_simple_table(
        doc,
        headers=["Activité", "LIRIE", "Partenaire"],
        rows=[
            [
                "Portefeuille propre",
                "Sous-traitant technique de la Plateforme",
                "Responsable du traitement",
            ],
            [
                "Sécurité, authentification et journaux",
                "Responsable distinct pour la sécurité de la Plateforme",
                "Responsable distinct pour la gestion de ses utilisateurs, "
                "habilitations, appareils et identifiants",
            ],
            [
                "Mise en relation et traçabilité — Course transmise",
                "Responsable pour les finalités de la Plateforme",
                "Responsable indépendant pour l'évaluation et "
                "l'acceptation de la demande",
            ],
            [
                "Exécution du transport et facturation au client",
                "—",
                "Responsable du traitement",
            ],
            [
                "Géolocalisation (GPS) des chauffeurs",
                "Sous-traitant technique pour la collecte, l'hébergement et "
                "la mise à disposition ; responsable distinct pour la "
                "sécurité de la Plateforme, la défense de ses droits, sa "
                "propre facturation et les statistiques irréversiblement "
                "anonymisées",
                "Responsable du traitement, notamment en qualité "
                "d'employeur et pour l'organisation des transports",
            ],
            [
                "Facturation LIRIE–Partenaire",
                "Responsable du traitement pour l'émission, le suivi et le "
                "recouvrement de ses factures",
                "Responsable distinct pour sa comptabilité, le contrôle et "
                "la contestation des factures reçues",
            ],
        ],
    )
    _add_para(
        doc,
        "Les rôles d'une institution à l'origine d'une demande sont régis "
        "par le contrat conclu entre cette institution et LIRIE, et non par "
        "le présent contrat.",
    )

    _add_heading(
        doc, "C.1 bis Obligations de LIRIE en qualité de sous-traitant", level=2
    )
    _add_bullets(
        doc,
        [
            "Lorsqu'elle agit en qualité de sous-traitant, notamment pour "
            "l'hébergement et le traitement technique des données de "
            "missions du portefeuille propre, à l'exclusion des traitements "
            "de sécurité, d'authentification, de journalisation, de "
            "prévention des abus et d'administration de la Plateforme pour "
            "lesquels LIRIE agit comme responsable distinct, LIRIE traite "
            "les données exclusivement sur instructions documentées du "
            "responsable du traitement concerné, sans les réutiliser pour "
            "ses propres finalités ;",
            "LIRIE veille à ce que ses collaborateurs ayant accès aux "
            "données soient soumis à une obligation de confidentialité ;",
            "LIRIE met en œuvre des mesures techniques et organisationnelles "
            "appropriées (contrôles d'accès, journalisation, sauvegardes), "
            "proportionnées au risque ;",
            "LIRIE apporte une assistance raisonnable au responsable du "
            "traitement pour répondre aux demandes d'exercice des droits "
            "des personnes concernées ;",
            "LIRIE informe le responsable du traitement concerné des "
            "incidents de sécurité affectant les données dont elle a "
            "connaissance, avec les informations disponibles utiles à leur "
            "analyse ;",
            "LIRIE est autorisée de manière générale à recourir à des "
            "sous-traitants ultérieurs, sous réserve d'en informer "
            "préalablement le Partenaire et de lui permettre de formuler "
            "une objection motivée dans un délai raisonnable ;",
            "Tout transfert international de données s'effectue dans le "
            "respect des exigences de la LPD, avec des garanties "
            "appropriées lorsque celles-ci sont requises ;",
            "À la fin du contrat, les données sont restituées ou "
            "supprimées conformément à l'article C.6, sous réserve des "
            "obligations légales de conservation ;",
            "LIRIE se soumet à des audits raisonnables, sur demande "
            "motivée et moyennant un préavis suffisant, et met à "
            "disposition les informations nécessaires pour démontrer sa "
            "conformité aux obligations du présent article. Sauf incident "
            "grave ou exigence d'une autorité, un audit ne peut être "
            "demandé plus d'une fois par période de douze mois. Il est "
            "réalisé en priorité par examen documentaire ou à distance, "
            "pendant les heures ouvrables, par une personne indépendante "
            "soumise à la confidentialité. Il ne peut donner accès aux "
            "données d'autres clients, au code source ou aux informations "
            "dont la divulgation compromettrait la sécurité. Les frais "
            "sont supportés par le Partenaire, sauf non-conformité "
            "substantielle imputable à LIRIE.",
        ],
    )

    _add_heading(doc, "C.2 Géolocalisation (GPS)", level=2)
    _add_para(
        doc,
        "Le Partenaire agit comme responsable du traitement pour la "
        "géolocalisation de ses chauffeurs, notamment en sa qualité "
        "d'employeur et pour l'organisation des transports. LIRIE agit "
        "comme sous-traitant pour la collecte, l'hébergement et la mise à "
        "disposition technique des positions. LIRIE agit toutefois comme "
        "responsable distinct pour les traitements strictement nécessaires "
        "à la sécurité de sa Plateforme, à la défense de ses droits, à sa "
        "propre facturation et à la production de statistiques "
        "irréversiblement anonymisées.",
    )
    _add_para(
        doc,
        "Les données de géolocalisation collectées par la Plateforme sont "
        "traitées aux fins de sécurité des personnes, d'organisation "
        "opérationnelle des courses, de preuve d'exécution, ainsi que, "
        "lorsque LIRIE agit comme responsable distinct, de défense de ses "
        "droits, de facturation propre et de statistiques irréversiblement "
        "anonymisées ou agrégées de manière à ne permettre raisonnablement "
        "ni l'identification ni la réidentification d'une personne — à "
        "l'exclusion de toute finalité de surveillance comportementale du "
        "personnel du Partenaire.",
    )
    _add_para(
        doc,
        "Le Partenaire est responsable d'informer préalablement ses "
        "chauffeurs des modalités, finalités, destinataires et durées du "
        "traitement de leur géolocalisation.",
    )
    _add_para(
        doc,
        "Le Partenaire a accès aux données de géolocalisation de sa propre "
        "flotte. Le client authentifié peut, lorsque la fonctionnalité le "
        "permet, accéder aux données de géolocalisation relatives à sa propre "
        "course.",
    )

    _add_heading(doc, "C.3 Durées de conservation", level=2)
    retention_categories = (
        compliance.get("retention_categories") or RETENTION_CATEGORIES
    )
    _add_bullets(
        doc,
        [str(cat.get("description") or "") for cat in retention_categories],
    )

    _add_heading(doc, "C.4 Sous-traitants techniques actifs", level=2)
    providers = compliance.get("providers") or active_technical_providers()
    if providers:
        headers = [
            "Prestataire",
            "Service",
            "Rôle",
            "Catégories de données",
            "Pays ou région de traitement",
            "Garanties",
        ]
        rows = []
        for p in providers:
            rows.append(
                [
                    str(p.get("name") or "—"),
                    str(p.get("service") or "—"),
                    str(p.get("legal_role") or "—"),
                    ", ".join(p.get("data_categories") or []) or "—",
                    ", ".join(p.get("processing_countries") or []) or "—",
                    str(p.get("transfer_guarantees") or "—"),
                ]
            )
        _add_simple_table(doc, headers=headers, rows=rows)
    _add_para(
        doc,
        "Pour Google Maps Platform, LIRIE s'engage à ne pas transmettre dans "
        "les requêtes API le nom du patient, sa date de naissance, une note "
        "médicale, un numéro de téléphone ni une référence de dossier. Les "
        "requêtes se limitent aux adresses ou coordonnées strictement "
        "nécessaires au géocodage, aux itinéraires et à la cartographie.",
    )
    _add_para(
        doc,
        "LIRIE informe le Partenaire de tout changement significatif de "
        "sous-traitant technique afin de lui permettre de formuler une "
        "objection motivée dans un délai raisonnable.",
    )

    _add_heading(doc, "C.5 Gestion des incidents", level=2)
    _add_para(
        doc,
        "LIRIE dirige la gestion des incidents affectant l'infrastructure et "
        "les traitements sous sa responsabilité de plateforme. Chaque Partie "
        "notifie l'autre sans délai injustifié, et dans la mesure du possible "
        "dans les 24 heures, de tout incident de sécurité pertinent la "
        "concernant ; la Partie responsable du traitement concerné procède, le "
        "cas échéant, aux annonces légales requises.",
    )
    _add_para(
        doc,
        "Ce délai de 24 heures constitue un objectif contractuel entre les "
        "Parties ; il ne remplace ni l'analyse de risque exigée par la loi, "
        "ni les délais légaux de notification au Préposé fédéral à la "
        "protection des données et à la transparence (PFPDT) incombant au "
        "responsable du traitement concerné.",
    )

    _add_heading(doc, "C.6 Fin du contrat", level=2)
    _add_para(
        doc,
        "À la fin du contrat, les données sont restituées ou supprimées selon "
        "les durées de conservation fixées à l'article C.3, sous réserve des "
        "obligations légales de conservation applicables (notamment "
        "comptables et fiscales).",
    )

    # ——— SIGNATURES ———
    _section_page_break(doc, label="page 8 — signatures")
    _add_heading(doc, "SIGNATURES", level=1, space_before_pt=16)
    _add_para(
        doc,
        "Les Parties déclarent avoir reçu, lu et accepté les Parties A "
        "(contrat-cadre), B (Annexe financière) et C (Annexe de protection "
        "des données) du présent contrat, lesquelles en forment un tout "
        "indivisible.",
        keep_with_next=True,
    )
    _add_para(doc, "Fait à Genève", keep_with_next=True, space_after_pt=4)
    _add_para(
        doc,
        f"Référence : {reference}",
        keep_with_next=True,
        space_after_pt=12,
    )
    operator_signatory = (
        _operator_natural_person_name(operator)
        if _is_sole_proprietor(operator)
        else (operator.get("signatory_name") or _operator_natural_person_name(operator))
    )
    partner_co_signatory = None
    if signatory_authority.get(
        "signature_mode"
    ) == "collective" and signatory_authority.get("co_signatory_name"):
        partner_co_signatory = signatory_authority.get("co_signatory_name")
    _add_signatures_table(
        doc,
        operator_signatory=operator_signatory,
        partner_signatory=partner.get("signatory_name"),
        partner_co_signatory=partner_co_signatory,
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
