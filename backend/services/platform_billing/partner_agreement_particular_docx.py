"""Renderer DOCX interne du contrat particulier (même contenu que le PDF)."""

from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from services.platform_billing.partner_agreement_particular_content import (
    ParticularAgreementContent,
)

_LOGO_CANDIDATES = (
    Path(__file__).resolve().parents[2] / "assets" / "lirie" / "logo-lirie.png",
    Path("/app/assets/lirie/logo-lirie.png"),
)


def _set_run(run, *, size: float = 9.5, bold: bool = False):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    r = run._element
    rpr = r.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")


def _para(
    doc: DocxDocument,
    text: str,
    *,
    size: float = 9.5,
    bold: bool = False,
    center: bool = False,
    space_after: float = 2,
    space_before: float = 0,
    keep_with_next: bool = False,
):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.05
    if keep_with_next:
        p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    _set_run(run, size=size, bold=bold)
    return p


def _heading(doc: DocxDocument, text: str) -> None:
    _para(
        doc,
        text,
        size=10.5,
        bold=True,
        space_before=8,
        space_after=3,
        keep_with_next=True,
    )


def _page_break(doc: DocxDocument) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._element.append(br)


def build_particular_docx_bytes(content: ParticularAgreementContent) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    for path in _LOGO_CANDIDATES:
        if path.is_file():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            p.add_run().add_picture(str(path), width=Cm(1.6))
            break

    _para(doc, content.title, size=13, bold=True, center=True)
    _para(doc, content.subtitle, size=9, center=True)
    _para(doc, f"Référence : {content.reference}", size=9, center=True)
    _para(
        doc,
        f"Date d'effet commerciale : {content.effective_date_fr}"
        f"  ·  Version : {content.particular_version}",
        size=9,
        center=True,
        space_after=6,
    )
    _para(doc, content.pack_note, size=9, space_after=6)

    _heading(doc, "Article 1 — Identification des Parties")
    for col in content.parties:
        _para(doc, col.header, size=9.5, bold=True, space_before=3, space_after=2)
        for line in col.lines:
            _para(doc, line, size=9)

    _heading(doc, "Article 2 — Conditions particulières")
    for row in content.commercial_terms:
        _para(doc, f"{row.label} : {row.value}", size=9)

    if content.special_conditions:
        _heading(doc, "Conditions particulières complémentaires")
        for line in content.special_conditions:
            _para(doc, f"— {line}", size=9)

    _heading(doc, content.key_principles_title)
    for i, principle in enumerate(content.key_principles, start=1):
        _para(doc, f"{i}. {principle}", size=9)

    _page_break(doc)
    _heading(doc, "Article 4 — Clauses générales essentielles")
    _para(doc, content.clauses_intro, size=9, space_after=4)
    for clause in content.clauses:
        _para(
            doc,
            clause.title,
            size=9.5,
            bold=True,
            keep_with_next=True,
            space_before=4,
        )
        _para(doc, clause.body, size=9)

    _page_break(doc)
    _heading(doc, "Article 5 — Protection des données")
    for role in content.data_protection_roles:
        _para(doc, f"{role.treatment} : {role.role}", size=9)
    _para(doc, content.data_protection_summary, size=9, space_before=4)
    _para(doc, content.gps_summary, size=9)
    _para(doc, content.providers_summary, size=9)

    _heading(doc, "Article 6 — Documents contractuels incorporés")
    _para(
        doc,
        "Le Partenaire reconnaît avoir reçu et accepté, avant la signature :",
        size=9,
    )
    for doc_item in content.incorporated_documents:
        _para(doc, f"• {doc_item.label}", size=9)
        _para(doc, f"  Version : {doc_item.version}", size=9)
    _para(doc, content.acceptance_clause, size=9, space_before=4)

    _heading(doc, "Article 7 — Signatures")
    _para(doc, content.signature_intro, size=9, space_after=4)
    for sig in content.signatures:
        _para(doc, sig.side, size=9.5, bold=True, space_before=6)
        _para(doc, sig.name, size=9)
        _para(doc, sig.title, size=9)
        if sig.power_attestation:
            _para(doc, sig.power_attestation, size=8)
        _para(doc, "Lieu et date : ____________________", size=9, space_before=4)
        _para(doc, "Signature : ____________________", size=9, space_before=6)
        if sig.co_signatory_name:
            co = sig.co_signatory_name
            if sig.co_signatory_title:
                co = f"{co}, {sig.co_signatory_title}"
            _para(doc, f"Co-signataire : {co}", size=9, space_before=4)
            _para(doc, "Signature : ____________________", size=9, space_before=4)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
